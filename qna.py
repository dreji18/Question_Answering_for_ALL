# -*- coding: utf-8 -*-
"""
Created on Mon Feb  8 18:04:27 2021

@author: rejid4996
"""

# loading the packages
import docx
import os
import urllib.request
import streamlit as st
import re 
import wikipedia
from rake_nltk import Rake
import torch
import streamlit as st 
from transformers import DistilBertTokenizer, DistilBertForQuestionAnswering
from rank_bm25 import BM25Okapi
from functools import reduce
import operator
from youtubesearchpython import VideosSearch
import sys
from io import BytesIO
import requests
from docx import Document
from docx.shared import Inches
import random
from io import BytesIO
from pathlib import Path
from PIL import Image
import tweepy as tw
import openai

@st.cache
# keyword extraction function
def keyword_extractor(query):
    """
    Rake has some features:
        1. convert automatically to lower case
        2. extract important key phrases
        3. it will extract combine words also (eg. Deep Learning, Capital City)
    """
    r = Rake() # Uses stopwords for english from NLTK, and all puntuation characters.
    r.extract_keywords_from_text(query)
    keywords = r.get_ranked_phrases() # To get keyword phrases ranked highest to lowest.
    return keywords

@st.cache
# data collection using wikepedia
def data_collection(search_words):
    """wikipedia"""
    search_query = ' '.join(search_words)
    wiki_pages = wikipedia.search(search_query, results = 5)
    
    information_list = []
    pages_list = []
    for i in wiki_pages:
        try:
            info = wikipedia.summary(i)
            if any(word in info.lower() for word in search_words):
                information_list.append(info)
                pages_list.append(i)
        except:
            pass
    
    original_info = information_list
    information_list = [item[:1000] for item in information_list] # limiting the word len to 512
    
    return information_list, pages_list, original_info

@st.cache
# document ranking function
def document_ranking(documents, query, n):
    """BM25"""
    try:
        tokenized_corpus = [doc.split(" ") for doc in documents]
        bm25 = BM25Okapi(tokenized_corpus)
        tokenized_query = query.split(" ")
        doc_scores = bm25.get_scores(tokenized_query)
        datastore = bm25.get_top_n(tokenized_query, documents, n)
    except:
        pass
    return datastore

@st.cache
def answergen(context, question):
    """DistilBert"""
    tokenizer = DistilBertTokenizer.from_pretrained('distilbert-base-uncased',return_token_type_ids = True)
    model = DistilBertForQuestionAnswering.from_pretrained('distilbert-base-uncased-distilled-squad')
    encoding = tokenizer.encode_plus(question, context)
    input_ids, attention_mask = encoding["input_ids"], encoding["attention_mask"]
    start_scores, end_scores = model(torch.tensor([input_ids]), attention_mask=torch.tensor([attention_mask]))
    ans_tokens = input_ids[torch.argmax(start_scores) : torch.argmax(end_scores)+1]
    answer_tokens = tokenizer.convert_ids_to_tokens(ans_tokens , skip_special_tokens=True)
    answer_tokens_to_string = tokenizer.convert_tokens_to_string(answer_tokens)

    return answer_tokens_to_string

@st.cache
def extract_images(pages, keyword_list):
    image_names_list = []
    for i in range(len(pages)):
        page = wikipedia.page(pages[i])
        image_names = page.images
        image_names_list.append(image_names)
    
    image_names_list = reduce(operator.concat, image_names_list)
           
    final_images = []
    for i in keyword_list:
        for j in image_names_list:   
            if re.sub('[^A-Za-z0-9]+', '', i) in re.sub('[^A-Za-z0-9]+', '', j).lower():
                try:
                    if urllib.request.urlopen(j).length > 10000:
                        final_images.append(j)
                except:
                    final_images.append(j)
                    pass
                    
    return list(set(final_images))

@st.cache
def extract_images_alternate(pages, keyword_list):    
    images_list = []
    for i in range(len(pages)):
        page = wikipedia.page(pages[i])
        try:
            if urllib.request.urlopen(page.images[0]).length > 10000:
                images_list.append(page.images[0])
        except:
            pass
    return list(set(images_list))

@st.cache
# extract videos
def extract_video(keyword_list):
    stopwords = ['what', 'who', 'whose', 'why', 'where', 'is']
    filtered = []
    for query in keyword_list:
        querywords = query.split()
        resultwords  = [word for word in querywords if word.lower() not in stopwords]
        filtered.append(' '.join(resultwords))
    
    youtube_search = ' '.join(filtered)

    result_list = []
    
    #for word in youtube_search:
    videosSearch = VideosSearch(youtube_search, limit = 2)
    videos_dict = videosSearch.result()
    result_list.append(videos_dict)
    
    videos_list = []
    for i in result_list:
        videos_list.append(i['result'][0]['link'])
        videos_list.append(i['result'][1]['link'])
    
    return videos_list, youtube_search

def tweets_extraction(search_words):
    # access permission
#    access_token =" "
#    access_token_secret =" "
#    consumer_key=" "
#    consumer_secret=" "
    
    auth = tw.OAuthHandler(consumer_key, consumer_secret)
    auth.set_access_token(access_token, access_token_secret)
    api = tw.API(auth, wait_on_rate_limit=True)
    
    date_since = "2020-11-16" # recent information
    
    # Collect tweets
    search_words = '+'.join([i.replace(' ', '') for i in search_words])
    new_search = search_words + " -filter:retweets"
    tweets = tw.Cursor(api.search,
                  q=new_search,
                  lang="en",
                  since=date_since).items(5)
    
    # Collect a list of tweets
    twitter_data = [tweet.text for tweet in tweets]
    
    return twitter_data

# contect creation function
def content_creation(images, keyword_list, original_data, query):
    doc = docx.Document() 
    doc.add_heading(query.title(), 0) 
    
    if len(images) > 3:
        image_id = random.sample(range(0, len(images)), 3)
        
        for i in [0, 1, 2]:
            doc_para = doc.add_paragraph(original_data[i])
            try:
                response = requests.get(images[image_id[i]])
                binary_img = BytesIO(response.content) 
                doc.add_picture(binary_img, width=Inches(3))
            except:
                pass
    else:
        for i in range(0, len(original_data)):
            doc_para = doc.add_paragraph(original_data[i])
            try:
                response = requests.get(images[i])
                binary_img = BytesIO(response.content) 
                doc.add_picture(binary_img, width=Inches(3))
            except:
                pass
    
    footer_section = doc.sections[0]
    footer = footer_section.footer 
    
    footer_text = footer.paragraphs[0]
    footer_text.text = 'This article uses material from the Wikipedia article Metasyntactic variable, which is released under the Creative Commons Attribution-ShareAlike 3.0 Unported License'
    
    path_to_download_folder = str(os.path.join(Path.home(), "Downloads"))
    doc.save(path_to_download_folder + '\\'+ query.title() +'.docx')
    
    return doc 

# main function
def main():
    """NLP App with Streamlit"""
    st.title("Question Answering for All!!!") #checkbox subheader
    
    st.sidebar.success("Please reach out to https://www.linkedin.com/in/deepak-john-reji/ for more queries")
    st.sidebar.subheader("Data extraction using NLP model ")
    
    st.info("For more contents subscribe to my Youtube Channel https://www.youtube.com/channel/UCgOwsx5injeaB_TKGsVD5GQ")
    
    wallpaper = Image.open('D 4 Data.jpg')
    wallpaper = wallpaper.resize((700,350))
    st.image(wallpaper)
    
    search_string = st.sidebar.text_input("type your question here and press enter", "")

    #if st.button("Extract"):# creates a button named "Extract"
    try:
        keyword_list = keyword_extractor(search_string)
        if len(keyword_list) == 0:
            st.stop()
    except:
        st.stop()
        
    try:       
        information, pages, original_data = data_collection(keyword_list)
        if len(information) == 0:
            st.info("Sorry, No Wikipedia Results")
    except:
        st.write("Sorry! Can't extract any information")
    
    try:
        #openai.api_key = "xxxxxxxxxxxxxxxxxxxxxxxxxx"

        start_sequence = "\nA:"
        restart_sequence = "\n\nQ: "
        
        response = openai.Completion.create(
          engine="davinci",
          #prompt="\nQ:who is billie eilish?\nA:",
          prompt="\nQ:"+search_string+"\nA:",
          temperature=0,
          max_tokens=100,
          top_p=1,
          frequency_penalty=0,
          presence_penalty=0,
          stop=["\n"]
        )
        
        openai_answer = response['choices'][0]['text']
    except:
        openai_answer = " "
    
    #st.subheader("open ai answer")
    #st.success(openai_answer)


    
    try:
        answers_list = []
        datastore = document_ranking(information, search_string, 3)
        ordered_data = document_ranking(original_data, search_string, 3)
                         
        for i in range(len(datastore)):
            result = answergen(datastore[i], search_string)
            answers_list.append(result)
    
    except:
        datastore = []
        print("")
    
    try:
        videos, query = extract_video(keyword_list) # extraction of videos
    except:
        st.write("Sorry! No Video got extracted")
    
    try:
        images = extract_images(pages, keyword_list) # extraction of images
        if len(images) == 0: #if main images are not loaded
            images = extract_images_alternate(pages, keyword_list)
    except:
        print("")
    
    try:
        tweets_box = st.sidebar.checkbox("Click here to see some tweets")
        if tweets_box:
            tweet_data = tweets_extraction(keyword_list)
            if tweet_data != []:
                st.subheader("🎲 Tweets")
                for tweet in tweet_data:
                     st.markdown(tweet)
    except:
        print("")
    
    cols = st.beta_columns(2)
    cols[0].subheader("🎲 Open AI Answers")
    cols[1].subheader("🎲 DistilBERT Answers")
    
    # printing answers
    cols = st.beta_columns(2)    
    cols[0].success(openai_answer)
    if answers_list != []:
        answers = '\n\n'.join(answers_list)
        if answers != '':
            cols[1].info(answers)

    if datastore != []:        
        st.subheader("🎲 Context")
        datastore = '\n\n'.join(datastore)
        st.markdown(datastore)
    
    try:
        if len(videos)!= 0:         
            st.subheader("🎲 Videos")        
            for i in videos:
                st.video(i)
    except:
        st.write("Sorry! No Videos unable to load")
    
    try:
        if len(images)!= 0:
            st.subheader("🎲 Images")  
            cols = st.beta_columns(2)
            num = 0
            for i in images:
                if i[-4:] == ".ogg":
                    st.markdown("Ha! I found some Audio for you")
                    st.audio(i, format='audio/ogg')
                else:
                    if num == 0:
                        cols[0].image(
                            i, width=None, # Manually Adjust the width of the image as per requirement,
                            use_column_width=True, caption=i
                        )
                        num = 1
                    else:
                        cols[1].image(
                            i, width=None, # Manually Adjust the width of the image as per requirement,
                            use_column_width=True, caption=i
                        )
                        num = 0
    except:
        st.write("Sorry! Images unable to load")
    
    if original_data != []:
        download_contents = st.sidebar.checkbox("Download the Contents") 
        if download_contents:      
            content = content_creation(images, keyword_list, original_data, query)
                
    st.info("All the data displayed in this app are from Wikipedia, Twitter and Youtube. This article uses material from the Wikipedia article Metasyntactic variable, which is released under the Creative Commons Attribution-ShareAlike 3.0 Unported License")
                   
# calling the main function
if __name__ == "__main__":
    main()
